"""
Copy the well-formatted RevolvingDoor_Sample.docx from F:\Github
and apply text changes (timer-based logic) while preserving formatting.
"""
import shutil, copy
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

SRC = r'F:\Github\RevolvingDoor_Sample.docx'
DST = r'f:\Github\AutomationBullShit\RevolvingDoor_Sample.docx'

# Step 1: Copy good file
shutil.copy2(SRC, DST)

doc = Document(DST)
body = doc.element.body

# ── helpers ──────────────────────────────────────────────────────────
def para_text(p):
    return p.text.strip()

def find_para(substr):
    """Find first paragraph containing substr."""
    for i, p in enumerate(doc.paragraphs):
        if substr in p.text:
            return i, p
    return None, None

def replace_para_text(p, new_text):
    """Replace all text in paragraph with new_text, keeping first run's formatting."""
    runs = p.runs
    if not runs:
        return
    # Keep first run formatting, set its text, clear rest
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''

def replace_in_runs(p, old, new):
    """Simple text replacement across paragraph (joining all runs, then redistributing)."""
    full = p.text
    if old not in full:
        return False
    new_full = full.replace(old, new)
    runs = p.runs
    if not runs:
        return False
    runs[0].text = new_full
    for r in runs[1:]:
        r.text = ''
    return True

def delete_paragraph(p):
    """Remove paragraph element from body."""
    elem = p._element
    elem.getparent().remove(elem)

def clone_para_with_text(template_para, new_text):
    """Clone a paragraph XML element and replace its text."""
    new_elem = copy.deepcopy(template_para._element)
    # Clear runs and set text
    for r in new_elem.findall(qn('w:r')):
        for t in r.findall(qn('w:t')):
            t.text = ''
    # Set first run text
    first_r = new_elem.find(qn('w:r'))
    if first_r is not None:
        t_elem = first_r.find(qn('w:t'))
        if t_elem is not None:
            t_elem.text = new_text
            t_elem.set(qn('xml:space'), 'preserve')
    return new_elem

# ── 1. Title: add period ─────────────────────────────────────────────
i, p = find_para('Вращающаяся дверь')
if p and not p.text.strip().endswith('.'):
    # Add period+space to title
    runs = p.runs
    if runs:
        last_run = runs[-1]
        if not last_run.text.rstrip().endswith('.'):
            last_run.text = last_run.text.rstrip() + '. '

# ── 2. Add "секция – Section" line after "Описание задачи." ──────────
# The broken file has this line; good file doesn't. Let's add it.
i_desc, p_desc = find_para('Описание задачи.')
if i_desc is not None:
    # Find the next non-empty paragraph (the description) and insert before it
    desc_para = doc.paragraphs[i_desc]
    # Clone formatting from description paragraph
    new_elem = copy.deepcopy(desc_para._element)
    for r in new_elem.findall(qn('w:r')):
        for t in r.findall(qn('w:t')):
            t.text = ''
    first_r = new_elem.find(qn('w:r'))
    if first_r is not None:
        # Remove bold if present
        rpr = first_r.find(qn('w:rPr'))
        if rpr is not None:
            b = rpr.find(qn('w:b'))
            if b is not None:
                rpr.remove(b)
        t_elem = first_r.find(qn('w:t'))
        if t_elem is not None:
            t_elem.text = 'секция – Section (A, B, C)'
            t_elem.set(qn('xml:space'), 'preserve')
    # Insert after empty paragraph following "Описание задачи."
    # Find the right position - after desc_para, add empty + new line
    desc_elem = desc_para._element
    next_sib = desc_elem.getnext()
    if next_sib is not None:
        # Insert before the next sibling (which should be empty para or the description)
        next_sib.addprevious(new_elem)
        # Add empty paragraph before it
        empty = copy.deepcopy(desc_para._element)
        for r in empty.findall(qn('w:r')):
            empty.remove(r)
        new_elem.addprevious(empty)

# ── 3. Change main description paragraph (P2 in good) ───────────────
i, p = find_para('Вращающаяся дверь — это трёхсекционная')
if p:
    # Need to rebuild runs. The good file has runs with bold variable names.
    # Strategy: clear all runs, rebuild with correct text.
    # First, gather formatting from existing runs
    runs = p.runs
    if runs:
        # Get the font size and style from first run
        fmt_size = runs[0].font.size

        # Clear all existing runs
        for r in runs:
            r_elem = r._element
            r_elem.getparent().remove(r_elem)

        # Add new runs with correct text
        from docx.shared import Pt
        SZ14 = Pt(14)

        def add_r(para, text, bold=False):
            r = para.add_run(text)
            r.font.size = fmt_size or SZ14
            if bold:
                r.font.bold = True
            return r

        add_r(p, 'Вращающаяся дверь — это трёхсекционная (')
        add_r(p, 'SECTION_COUNT', bold=True)
        add_r(p, ' = 3) дверь, установленная в проёме здания, обеспечивающая проход людей с двух сторон (')
        add_r(p, 'SideA / SideB', bold=True)
        add_r(p, '). В нормальном (исходном) состоянии дверь неподвижна. '
              'При приближении пользователя к одной из сторон (датчики ')
        add_r(p, 'presenceSideA / presenceSideB', bold=True)
        add_r(p, ') двигатель (')
        add_r(p, 'motorOn', bold=True)
        add_r(p, ') запускает вращение двери с заданной скоростью (')
        add_r(p, 'ROTATION_SPEED', bold=True)
        add_r(p, '). Вращение продолжается в течение заданного времени (')
        add_r(p, 'ROTATION_TIME', bold=True)
        add_r(p, '). Если за время вращения повторно срабатывает один из сенсоров, '
              'таймер сбрасывается и отсчёт начинается заново. '
              'Вращение останавливается по истечении таймера. '
              'При оказании давления на секционные перегородки (датчик ')
        add_r(p, 'partitionPressure', bold=True)
        add_r(p, ') вращение приостанавливается на время ')
        add_r(p, 'PAUSE_DURATION', bold=True)
        add_r(p, ', после чего возобновляется, если таймер ещё не истёк.')

# ── 4. Change speed paragraph (P4 in good) ──────────────────────────
i, p = find_para('Скорость вращения двери регулируется')
if p:
    runs = p.runs
    if runs:
        fmt_size = runs[0].font.size
        for r in runs:
            r._element.getparent().remove(r._element)
        from docx.shared import Pt
        SZ14 = Pt(14)
        def add_r2(para, text, bold=False):
            r = para.add_run(text)
            r.font.size = fmt_size or SZ14
            if bold:
                r.font.bold = True
        add_r2(p, 'Дверь вращается с постоянной скоростью (')
        add_r2(p, 'ROTATION_SPEED', bold=True)
        add_r2(p, '). Время вращения после последнего срабатывания сенсора задаётся константой ')
        add_r2(p, 'ROTATION_TIME', bold=True)
        add_r2(p, '.')

# ── 5. Remove "Оператор может имитировать нахождение пользователя в зоне вращения" ──
i, p = find_para('Оператор может имитировать нахождение пользователя в зоне вращения')
if p:
    delete_paragraph(p)

# ── 6. Change controller description lines ──────────────────────────
i, p = find_para('Двигатель запускается при обнаружении пользователя датчиками приближения')
if p:
    replace_para_text(p, 'Двигатель запускается при срабатывании одного из двух сенсоров движения.')

i, p = find_para('Двигатель останавливается, когда зона вращения пуста')
if p:
    replace_para_text(p, 'Двигатель останавливается по истечении таймера ROTATION_TIME.')

# ── 7. Environment: remove EnterDoor, LeaveDoor ─────────────────────
i, p = find_para('Environment:')
if p:
    replace_in_runs(p, 'EnterDoor, LeaveDoor, ', '')
    # Also remove "EnterDoor, LeaveDoor, " if it appears differently
    replace_in_runs(p, ', EnterDoor, LeaveDoor', '')

# ── 8. Constants: SAFE_SPEED → ROTATION_TIME ────────────────────────
for i, p in enumerate(doc.paragraphs):
    if 'SAFE_SPEED' in p.text:
        replace_in_runs(p, 'SAFE_SPEED', 'ROTATION_TIME')

# ── 9. Controls: remove userInDoor ───────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if 'userInDoor' in p.text and 'Controls' in p.text:
        replace_in_runs(p, ', userInDoor', '')
        replace_in_runs(p, 'userInDoor, ', '')

# ── 10. Remove "Имитировать вход пользователя в зону вращения" ───────
i, p = find_para('Имитировать вход пользователя в зону вращения')
if p:
    delete_paragraph(p)

# ── 11. Change Controls section under "Ручной режим" ────────────────
i, p = find_para('Имитировать вход/выход пользователя в зону вращения')
if p:
    replace_para_text(p, 'При срабатывании сенсора дверь вращается ROTATION_TIME. При закрытых индикаторах — ошибка.')

# ── 12. VAR_GLOBAL: remove userInDoor lines ─────────────────────────
# Find and remove "(* User in rotation zone *)" and "userInDoor  : BOOL;"
i, p = find_para('(* User in rotation zone *)')
if p:
    delete_paragraph(p)

i, p = find_para('userInDoor')
while p is not None:
    if 'userInDoor' in p.text and ':=' not in p.text and 'Controls' not in p.text and 'IF' not in p.text and 'NOT' not in p.text:
        # This is a VAR declaration line
        if ': BOOL' in p.text or 'BOOL;' in p.text:
            delete_paragraph(p)
    i, p = find_para('userInDoor')
    # Prevent infinite loop
    break

# More targeted: find exact "userInDoor  : BOOL;" paragraph
for p in list(doc.paragraphs):
    txt = p.text.strip()
    if txt == 'userInDoor  : BOOL;' or txt == 'userInDoor : BOOL;':
        delete_paragraph(p)

# ── 13. Plant Init: remove "userInDoor := FALSE;" ───────────────────
for p in list(doc.paragraphs):
    if p.text.strip() == 'userInDoor := FALSE;':
        delete_paragraph(p)

# ── 14. Controller MainControl: replace with timer-based logic ──────
# Find MainControl process and replace its states
# Strategy: find STATE idle, STATE rotating, STATE paused and replace their content

# First, let's handle the MainControl PROCESS declaration
# Add VAR CONSTANT and VAR sections after "PROCESS MainControl"
i_mc, p_mc = find_para('PROCESS MainControl')
if p_mc is not None:
    # The good file goes directly to "STATE idle" after "PROCESS MainControl"
    # We need to insert VAR CONSTANT / ROTATION_TIME / END_VAR / VAR / timer / END_VAR

    # Find "STATE idle" paragraph
    i_idle = None
    for j, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'STATE idle':
            i_idle = j
            break

    if i_idle is not None:
        idle_elem = doc.paragraphs[i_idle]._element

        # Create new paragraphs by cloning the STATE idle paragraph format
        lines_to_insert = [
            '  VAR CONSTANT',
            '    ROTATION_TIME : INT := 50;',
            '  END_VAR',
            '  VAR',
            '    timer : INT := 0;',
            '  END_VAR',
        ]

        for line in reversed(lines_to_insert):
            new_elem = copy.deepcopy(idle_elem)
            # Set text in first run
            for r in new_elem.findall(qn('w:r')):
                for t in r.findall(qn('w:t')):
                    t.text = ''
            first_r = new_elem.find(qn('w:r'))
            if first_r is not None:
                t_elem = first_r.find(qn('w:t'))
                if t_elem is not None:
                    t_elem.text = line
                    t_elem.set(qn('xml:space'), 'preserve')
                # Set bold for VAR/END_VAR lines
                rpr = first_r.find(qn('w:rPr'))
                if rpr is not None:
                    b = rpr.find(qn('w:b'))
                    if line.strip().startswith('VAR') or line.strip().startswith('END_VAR'):
                        if b is None:
                            etree.SubElement(rpr, qn('w:b'))
                    else:
                        if b is not None:
                            rpr.remove(b)
            idle_elem.addprevious(new_elem)

# ── 15. Change "STATE rotating" to "STATE rotating LOOPED" ──────────
i, p = find_para('STATE rotating')
if p and 'LOOPED' not in p.text:
    runs = p.runs
    if runs:
        runs[0].text = runs[0].text.replace('STATE rotating', 'STATE rotating LOOPED')

# ── 16. Replace STATE idle content ───────────────────────────────────
# Good file idle: IF → START PROCESS StartRotation; SET STATE rotating; END_IF
# Broken file idle: same + "timer := ROTATION_TIME;" after START PROCESS
# Find "START PROCESS StartRotation;" inside idle state and add timer line after it

# Find the sequence within idle state
for j, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'START PROCESS StartRotation;':
        # Check if next line is "SET STATE rotating;" (idle state context)
        if j + 1 < len(doc.paragraphs) and 'SET STATE rotating' in doc.paragraphs[j+1].text:
            # Insert "timer := ROTATION_TIME;" between these two
            set_state_elem = doc.paragraphs[j+1]._element
            new_elem = copy.deepcopy(p._element)
            for r in new_elem.findall(qn('w:r')):
                for t in r.findall(qn('w:t')):
                    t.text = ''
            first_r = new_elem.find(qn('w:r'))
            if first_r is not None:
                t_elem = first_r.find(qn('w:t'))
                if t_elem is not None:
                    t_elem.text = '      timer := ROTATION_TIME;'
                    t_elem.set(qn('xml:space'), 'preserve')
                # Remove bold
                rpr = first_r.find(qn('w:rPr'))
                if rpr is not None:
                    b = rpr.find(qn('w:b'))
                    if b is not None:
                        rpr.remove(b)
            set_state_elem.addprevious(new_elem)
            break

# ── 17. Replace STATE rotating content ──────────────────────────────
# Good file rotating state:
#   IF partitionPressure THEN ... END_IF
#   IF (NOT userInDoor AND NOT presenceSideA AND NOT presenceSideB) THEN ... StopRotation ... idle ... END_IF
#
# Broken file rotating state (timer-based):
#   IF (presenceSideA OR presenceSideB) THEN timer := ROTATION_TIME; END_IF
#   IF partitionPressure THEN ... PauseRotation ... paused ... END_IF
#   timer := timer - 1;
#   IF timer <= 0 THEN ... StopRotation ... idle ... END_IF

# Find the rotating state and replace the userInDoor check with timer logic
# Find "IF (NOT userInDoor" line
for j, p in enumerate(doc.paragraphs):
    if 'NOT userInDoor AND NOT presenceSideA AND NOT presenceSideB' in p.text:
        # This is the stop condition in rotating state. Replace with timer logic.
        # We need to:
        # 1. Add "IF (presenceSideA OR presenceSideB) THEN" + "timer := ROTATION_TIME;" + "END_IF" BEFORE the partitionPressure check
        # 2. Replace this line with "timer := timer - 1;"
        # 3. Change the next lines to timer-based stop

        # Replace current line with timer decrement
        replace_para_text(p, '    timer := timer - 1;')

        # Replace "START PROCESS StopRotation;" context - find it after this line
        # The next lines should be: START PROCESS StopRotation; SET STATE idle; END_IF
        # Change the IF condition
        # Actually we need to change the IF line before StopRotation
        # Let's find it: it should be right after timer := timer - 1
        if j + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[j + 1]
            if 'START PROCESS StopRotation' in next_p.text:
                # Need to add "IF timer <= 0 THEN" before this
                stop_elem = next_p._element
                new_elem = copy.deepcopy(stop_elem)
                for r in new_elem.findall(qn('w:r')):
                    for t in r.findall(qn('w:t')):
                        t.text = ''
                first_r = new_elem.find(qn('w:r'))
                if first_r is not None:
                    t_elem = first_r.find(qn('w:t'))
                    if t_elem is not None:
                        t_elem.text = '    IF timer <= 0 THEN'
                        t_elem.set(qn('xml:space'), 'preserve')
                stop_elem.addprevious(new_elem)
        break

# Now add the timer reset block before partitionPressure check in rotating state
# Find "IF partitionPressure THEN" that's inside rotating state
partition_indices = []
for j, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'IF partitionPressure THEN':
        partition_indices.append(j)

# The first occurrence should be in rotating state
if partition_indices:
    j = partition_indices[0]
    part_elem = doc.paragraphs[j]._element

    # Insert timer reset block before partitionPressure check
    timer_lines = [
        '    IF (presenceSideA OR presenceSideB) THEN',
        '      timer := ROTATION_TIME;',
        '    END_IF',
    ]

    for line in reversed(timer_lines):
        new_elem = copy.deepcopy(part_elem)
        for r in new_elem.findall(qn('w:r')):
            for t in r.findall(qn('w:t')):
                t.text = ''
        first_r = new_elem.find(qn('w:r'))
        if first_r is not None:
            t_elem = first_r.find(qn('w:t'))
            if t_elem is not None:
                t_elem.text = line
                t_elem.set(qn('xml:space'), 'preserve')
            rpr = first_r.find(qn('w:rPr'))
            if rpr is not None:
                b = rpr.find(qn('w:b'))
                if b is not None:
                    rpr.remove(b)
        part_elem.addprevious(new_elem)

# ── 18. Replace STATE paused content ────────────────────────────────
# Good: IF userInDoor THEN ... ELSIF ...
# Broken: IF timer > 0 THEN ... ELSIF ...
for p in doc.paragraphs:
    if p.text.strip() == 'IF userInDoor THEN':
        replace_para_text(p, '      IF timer > 0 THEN')

# ── 19. Controller VAR section: remove userInDoor ────────────────────
# Remove "(* User in rotation zone *)" in Controller section
found_controller = False
for p in list(doc.paragraphs):
    if 'PROGRAM Controller' in p.text:
        found_controller = True
    if found_controller:
        if p.text.strip() == '(* User in rotation zone *)':
            delete_paragraph(p)
        if p.text.strip() in ('userInDoor  : BOOL;', 'userInDoor : BOOL;'):
            delete_paragraph(p)

# ── 20. Table: update requirements ──────────────────────────────────
from docx.shared import Pt
SZ10 = Pt(10)

def set_cell_text(cell, text_runs):
    """Set cell text with formatting. text_runs is list of (text, bold) tuples."""
    # Clear existing paragraphs content
    for para in cell.paragraphs:
        for r in para.runs:
            r.text = ''

    # Use first paragraph
    para = cell.paragraphs[0]
    # Remove all existing runs
    for r_elem in para._element.findall(qn('w:r')):
        para._element.remove(r_elem)

    # Add new runs
    for text, bold in text_runs:
        r = para.add_run(text)
        r.font.size = SZ10
        if bold:
            r.font.bold = True

table = doc.tables[0]

# Row 2 [2]: Change description and trigger
set_cell_text(table.cell(2, 0), [('Дверь должна начать вращаться при активации одного из двух сенсоров движения', False)])
set_cell_text(table.cell(2, 1), [('presenceSideA.RE OR presenceSideB.RE', False)])

# Row 3 [3]: Change description and trigger/reaction
set_cell_text(table.cell(3, 0), [('Дверь должна вращаться в течение заданного времени ROTATION_TIME после последней активации сенсора', False)])
set_cell_text(table.cell(3, 1), [('presenceSideA.RE OR presenceSideB.RE', False)])
set_cell_text(table.cell(3, 3), [('tau(#ROTATION_TIME)', False)])
set_cell_text(table.cell(3, 5), [('motorOn', False)])
set_cell_text(table.cell(3, 6), [('NOT ', True), ('motorOn', False)])

# Row 5 [5]: Change trigger
set_cell_text(table.cell(5, 1), [('partitionPressure.FE AND timer > 0', False)])
set_cell_text(table.cell(5, 0), [('После снятия давления вращение возобновляется через PAUSE_DURATION если таймер ещё не истёк', False)])

# Row 9 [9]: Change description and invariant
set_cell_text(table.cell(9, 0), [('Дверь не должна вращаться при истёкшем таймере и отсутствии активации сенсоров', False)])
set_cell_text(table.cell(9, 5), [('NOT ', True), ('(motorOn AND timer <= 0)', False)])

# ── Save ─────────────────────────────────────────────────────────────
doc.save(DST)
print('Done! Saved to', DST)
