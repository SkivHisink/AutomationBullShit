# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Emu, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

SZ14 = Pt(14)  # 177800 EMU - body text size in original
SZ10 = Pt(10)  # 127000 EMU - table text size in original


def make_para(doc, alignment=None):
    """Create an empty paragraph with optional alignment."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.paragraph_format.alignment = alignment
    return p


def add_run(para, text, bold=False, size=SZ14):
    """Add a run to a paragraph with formatting matching original."""
    r = para.add_run(text)
    r.font.size = size
    if bold:
        r.font.bold = True
    return r


def add_simple(doc, text, bold=False, align=None):
    """Add a single-run paragraph."""
    p = make_para(doc, align)
    add_run(p, text, bold=bold)
    return p


def add_empty(doc):
    """Add an empty paragraph."""
    doc.add_paragraph()


def add_table_cell_text(cell, text, bold=False):
    """Set cell text with formatting matching original table."""
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = SZ10
    if bold:
        r.font.bold = True


# ================================================================
# Create new document and set up page to match original Sample.docx
# ================================================================
doc = Document()

# Page setup: A4 Landscape matching original
from docx.enum.section import WD_ORIENT
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width = Cm(29.7)
sec.page_height = Cm(21.0)
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.75)
sec.right_margin = Cm(1.21)

body = doc.element.body

# ============ TITLE ============
p = make_para(doc)
add_run(p, 'Вращающаяся дверь. ', bold=True)

add_empty(doc)

add_simple(doc, 'Описание задачи.', bold=True)

add_empty(doc)

add_simple(doc, 'секция \u2013 Section (A, B, C)')

add_empty(doc)

# ============ MAIN DESCRIPTION (JUSTIFY) ============
J = WD_ALIGN_PARAGRAPH.JUSTIFY

# Paragraph 1 - main description with bold variable names
p = make_para(doc, J)
add_run(p, 'Вращающаяся дверь \u2014 это трёхсекционная (')
add_run(p, 'SECTION_COUNT', bold=True)
add_run(p, ' = 3) дверь, установленная в проёме здания, '
           'обеспечивающая проход людей с двух сторон (')
add_run(p, 'SideA / SideB', bold=True)
add_run(p, '). В нормальном (исходном) состоянии дверь неподвижна. '
           'При приближении пользователя к одной из сторон (датчики ')
add_run(p, 'presenceSideA / presenceSideB', bold=True)
add_run(p, ') двигатель (')
add_run(p, 'motorOn', bold=True)
add_run(p, ') запускает вращение двери с заданной скоростью (')
add_run(p, 'ROTATION_SPEED', bold=True)
add_run(p, '). Вращение продолжается в течение заданного времени (')
add_run(p, 'ROTATION_TIME', bold=True)
add_run(p, '). Если за время вращения повторно срабатывает один из сенсоров, '
           'таймер сбрасывается и отсчёт начинается заново. '
           'Вращение останавливается по истечении таймера. '
           'При оказании давления на секционные перегородки (датчик ')
add_run(p, 'partitionPressure', bold=True)
add_run(p, ') вращение приостанавливается на время ')
add_run(p, 'PAUSE_DURATION', bold=True)
add_run(p, ', после чего возобновляется, если таймер ещё не истёк.')
add_run(p, ' Для ручного управления предусмотрены специальные кнопки ')
add_run(p, 'rotateForwardButton / rotateBackwardButton', bold=True)
add_run(p, ', позволяющие вращать дверь вперёд или назад, пока соответствующая кнопка удерживается нажатой.')

# Paragraph 2
p = make_para(doc, J)
add_run(p, 'Дверь оснащена датчиком угла поворота (')
add_run(p, 'doorAngle', bold=True)
add_run(p, '), позволяющим определить текущее положение секций. '
           'Для информирования пользователей о возможности входа используются световые индикаторы '
           'с каждой стороны (')
add_run(p, 'lightSideA / lightSideB', bold=True)
add_run(p, '): зелёный \u2014 проход разрешён, красный \u2014 проход запрещён '
           '(дверь заблокирована или приостановлена).')

# Paragraph 3
p = make_para(doc, J)
add_run(p, 'Дверь вращается с постоянной скоростью (')
add_run(p, 'ROTATION_SPEED', bold=True)
add_run(p, '). Время вращения после последнего срабатывания сенсора задаётся константой ')
add_run(p, 'ROTATION_TIME', bold=True)
add_run(p, '.')

add_empty(doc)
add_empty(doc)

# ============ MODEL OBJECT (environment control) ============
p = make_para(doc, J)
add_run(p, 'К модели объекта (управление со стороны среды)', bold=True)
add_run(p, ':')

add_simple(doc, 'Кнопка "Исходное".', align=J)
add_simple(doc, 'Оператор может имитировать подход пользователя с каждой стороны (кнопки ApproachSideA / ApproachSideB).', align=J)
add_simple(doc, 'Оператор может имитировать давление на перегородку (кнопка PressPartition / ReleasePartition).', align=J)
add_simple(doc, 'Оператор может нажимать специальные кнопки rotateForwardButton / rotateBackwardButton для ручного вращения двери вперёд или назад.', align=J)
add_simple(doc, 'Попытка войти при красном индикаторе приводит к выдаче сообщения об ошибке.', align=J)

add_empty(doc)

# ============ MODEL OBJECT (transform function) ============
p = make_para(doc, J)
add_run(p, 'К модели объекта (управление со стороны функции преобразования)', bold=True)
add_run(p, ':')

add_empty(doc)

add_simple(doc, 'Попытка запустить двигатель при активном давлении на перегородку приводит к выдаче сообщения об ошибке.', align=J)

add_empty(doc)

# ============ MODEL OBJECT (controller) ============
p = make_para(doc, J)
add_run(p, 'К модели объекта (управление со стороны контроллера)', bold=True)
add_run(p, ':')

add_simple(doc, 'Двигатель запускается при срабатывании одного из двух сенсоров движения.', align=J)
add_simple(doc, 'Двигатель останавливается по истечении таймера ROTATION_TIME.', align=J)
add_simple(doc, 'При давлении на перегородку двигатель останавливается на PAUSE_DURATION секунд.', align=J)
add_simple(doc, 'Светофоры переключаются в зависимости от состояния двери (вращение/остановка/пауза).', align=J)
add_simple(doc, 'При нажатии rotateForwardButton дверь должна вращаться вперёд, а при нажатии rotateBackwardButton — назад.', align=J)
add_simple(doc, 'Одновременное нажатие rotateForwardButton и rotateBackwardButton недопустимо и не должно запускать двигатель.', align=J)

add_empty(doc)
add_empty(doc)

# ============ CYBERPHYSICAL DIAGRAM ============
add_simple(doc, 'Киберфизическая диаграмма.', bold=True)

add_empty(doc)

p = make_para(doc, J)
add_run(p, 'Environment', bold=True)
add_run(p, ': ApproachSideA, ApproachSideB, LeaveSideA, LeaveSideB, PressPartition, ReleasePartition, PressForwardButton, ReleaseForwardButton, PressBackwardButton, ReleaseBackwardButton')

p = make_para(doc, J)
add_run(p, 'Constants', bold=True)
add_run(p, ': ROTATION_SPEED, ROTATION_TIME, PAUSE_DURATION, SECTION_COUNT')

p = make_para(doc, J)
add_run(p, 'Controls', bold=True)
add_run(p, ': ')

add_simple(doc, 'presenceSideA, presenceSideB, partitionPressure, rotateForwardButton, rotateBackwardButton', align=J)

p = make_para(doc, J)
add_run(p, 'Sensors: ', bold=True)

add_simple(doc, 'doorAngle, isRotating', align=J)

p = make_para(doc, J)
add_run(p, 'Actuators', bold=True)
add_run(p, ': motorOn, rotationForward')

p = make_para(doc, J)
add_run(p, 'Indicators: ', bold=True)
add_run(p, 'lightSideA, lightSideB')

add_empty(doc)
add_empty(doc)

# ============ MANUAL MODE ============
add_simple(doc, 'Ручной режим.')
add_simple(doc, 'Имитировать подход пользователя со стороны A, стороны B, убрать пользователя с A, убрать с B.')

add_empty(doc)

add_simple(doc, 'Имитировать давление на перегородку / снятие давления.')

add_empty(doc)

add_simple(doc, 'Нажать/отпустить rotateForwardButton для вращения двери вперёд.', align=J)
add_simple(doc, 'Нажать/отпустить rotateBackwardButton для вращения двери назад.', align=J)
add_simple(doc, 'При активном давлении на перегородку нажатие любой из кнопок ручного вращения должно приводить к ошибке.', align=J)
add_simple(doc, 'Одновременное нажатие rotateForwardButton и rotateBackwardButton недопустимо.', align=J)

add_empty(doc)

add_simple(doc, 'Установить индикаторы вручную.')

add_empty(doc)
add_empty(doc)

# ============ Controls (repeated) ============
add_simple(doc, 'Controls: ')
add_simple(doc, 'Имитировать подход пользователя со стороны A, стороны B.')
add_simple(doc, 'При срабатывании сенсора дверь вращается ROTATION_TIME. При закрытых индикаторах \u2014 ошибка.')
add_simple(doc, 'Нажатие rotateForwardButton / rotateBackwardButton запускает ручное вращение в выбранном направлении, пока кнопка удерживается нажатой.')

add_empty(doc)
add_empty(doc)
add_empty(doc)
add_empty(doc)
add_empty(doc)

# ============ OBJECT AND VISUALIZATION ============
add_simple(doc, 'Объект управления и визуализация', bold=True)

add_empty(doc)

p = make_para(doc, J)
add_run(p, 'Constants', bold=True)
add_run(p, ': ROTATION_SPEED, ROTATION_TIME, PAUSE_DURATION, SECTION_COUNT')

p = make_para(doc, J)
add_run(p, 'Controls', bold=True)
add_run(p, ': presenceSideA, presenceSideB, partitionPressure, rotateForwardButton, rotateBackwardButton')

p = make_para(doc, J)
add_run(p, 'Sensors: ', bold=True)
add_run(p, 'doorAngle, isRotating')

p = make_para(doc, J)
add_run(p, 'Actuators', bold=True)
add_run(p, ': motorOn, rotationForward')

p = make_para(doc, J)
add_run(p, 'Indicators: ', bold=True)
add_run(p, 'lightSideA, lightSideB')

add_empty(doc)
add_empty(doc)

# ============ VAR_GLOBAL (line-by-line like original) ============
# Each line is a separate paragraph, matching original format

def add_code_line(doc, text, bold=False):
    p = make_para(doc)
    add_run(p, text, bold=bold)
    return p

def add_var_line(doc, varname, vartype='BOOL'):
    """Add variable declaration: name normal, type bold (like original)."""
    p = make_para(doc)
    add_run(p, '         ' + varname + ' ')
    add_run(p, ' : ' + vartype + ';', bold=True)
    return p

add_code_line(doc, 'VAR_GLOBAL (* VAR_INPUT *)', bold=True)
add_code_line(doc, '\t(* Presence sensors *)', bold=True)
add_var_line(doc, 'presenceSideA')
add_var_line(doc, 'presenceSideB')
add_code_line(doc, '\t(* Partition pressure sensor *)', bold=True)
add_var_line(doc, 'partitionPressure')
add_code_line(doc, '\t(* Manual rotation buttons *)', bold=True)
add_var_line(doc, 'rotateForwardButton')
add_var_line(doc, 'rotateBackwardButton')
add_code_line(doc, '\t(* Door angle sensor *)', bold=True)
add_var_line(doc, 'doorAngle', 'REAL')
add_code_line(doc, '\t(* Rotation status *)', bold=True)
add_var_line(doc, 'isRotating')
add_code_line(doc, '(* END_VAR', bold=True)
add_code_line(doc, 'VAR *) (* VAR_OUTPUT *)', bold=True)
add_code_line(doc, '\t(* Motor control *)', bold=True)
add_var_line(doc, 'motorOn')
add_var_line(doc, 'rotationForward')
add_code_line(doc, '\t(* Traffic lights *)', bold=True)
add_var_line(doc, 'lightSideA')
add_var_line(doc, 'lightSideB')
add_code_line(doc, 'END_VAR', bold=True)

add_empty(doc)

# ============ PROGRAM Plant ============
plant_lines = [
    ('PROGRAM Plant (* RevolvingDoor *) ', True),
    ('(* Plant *)', True),
    ('PROCESS Init ', True),
    ('  STATE begin', False),
    ('    (* inputs: *)', True),
    ('    presenceSideA := FALSE;', False),
    ('    presenceSideB := FALSE;', False),
    ('    partitionPressure := FALSE;', False),
    ('    rotateForwardButton := FALSE;', False),
    ('    rotateBackwardButton := FALSE;', False),
    ('    doorAngle := 0.0;', False),
    ('    isRotating := FALSE;', False),
    ('(* outputs: *)', True),
    ('    motorOn := FALSE;', False),
    ('    rotationForward := TRUE;', False),
    ('    lightSideA := TRUE;', False),
    ('    lightSideB := TRUE;', False),
    ('    START PROCESS MotorSim;', False),
    ('    START PROCESS DoorRotationSim;', False),
    ('    START PROCESS PressureSim;', False),
    ('    STOP;', False),
    ('  END_STATE', False),
    ('END_PROCESS', False),
    ('', False),
    ('PROCESS MotorSim ', True),
    ('\tVAR CONSTANT', True),
    ('\t\tROTATION_SPEED : REAL := 2.0;', False),
    ('\tEND_VAR', True),
    ('\tSTATE check_motor LOOPED ', False),
    ('\t\tIF motorOn THEN', False),
    ('\t\t\tisRotating := TRUE;', False),
    ('\t\tELSE', False),
    ('\t\t\tisRotating := FALSE;', False),
    ('\t\tEND_IF', False),
    ('\tEND_STATE', False),
    ('END_PROCESS', False),
    ('', False),
    ('PROCESS DoorRotationSim', True),
    ('\tVAR CONSTANT', True),
    ('\t\tROTATION_SPEED : REAL := 2.0;', False),
    ('\t\tMAX_ANGLE : REAL := 360.0;', False),
    ('\tEND_VAR', True),
    ('\tSTATE check_rotation LOOPED ', False),
    ('\t\tIF isRotating THEN', False),
    ('\t\t\tIF rotationForward THEN', False),
    ('\t\t\t\tdoorAngle := doorAngle + ROTATION_SPEED;', False),
    ('\t\t\tELSE', False),
    ('\t\t\t\tdoorAngle := doorAngle - ROTATION_SPEED;', False),
    ('\t\t\tEND_IF', False),
    ('\t\tEND_IF', False),
    ('\t\tIF doorAngle >= MAX_ANGLE THEN ', False),
    ('\t\t\tdoorAngle := doorAngle - MAX_ANGLE;', False),
    ('\t\tEND_IF', False),
    ('\t\tIF doorAngle < 0.0 THEN ', False),
    ('\t\t\tdoorAngle := doorAngle + MAX_ANGLE;', False),
    ('\t\tEND_IF', False),
    ('\tEND_STATE', False),
    ('END_PROCESS', False),
    ('', False),
    ('PROCESS PressureSim', True),
    ('\tVAR CONSTANT', True),
    ('\t\tPAUSE_DURATION : REAL := 30.0;', False),
    ('\tEND_VAR', True),
    ('\tVAR', True),
    ('\t\tpauseCounter : REAL := 0.0;', False),
    ('\tEND_VAR', True),
    ('\tSTATE check_pressure LOOPED ', False),
    ('\t\tIF partitionPressure THEN', False),
    ('\t\t\tpauseCounter := PAUSE_DURATION;', False),
    ('\t\tEND_IF', False),
    ('\t\tIF pauseCounter > 0.0 THEN', False),
    ('\t\t\tpauseCounter := pauseCounter - 1.0;', False),
    ('\t\tEND_IF', False),
    ('\tEND_STATE', False),
    ('END_PROCESS', False),
    ('END_PROGRAM', True),
]

for text, bold in plant_lines:
    add_code_line(doc, text, bold=bold)

add_empty(doc)

# ============ PROGRAM Controller ============
ctrl_lines = [
    ('PROGRAM Controller', True),
    ('VAR (* VAR_INPUT *)', True),
    ('\t(* Presence sensors *)', True),
    ('         presenceSideA  : BOOL;', False),
    ('         presenceSideB : BOOL;', False),
    ('\t(* Partition pressure sensor *)', True),
    ('         partitionPressure : BOOL;', False),
    ('\t(* Manual rotation buttons *)', True),
    ('         rotateForwardButton : BOOL;', False),
    ('         rotateBackwardButton : BOOL;', False),
    ('\t(* Door angle sensor *)', True),
    ('         doorAngle : REAL;', False),
    ('\t(* Rotation status *)', True),
    ('         isRotating : BOOL;', False),
    ('END_VAR', True),
    ('VAR (* VAR_OUTPUT *)', True),
    ('\t(* Motor control *)', True),
    ('         motorOn : BOOL;', False),
    ('         rotationForward : BOOL;', False),
    ('\t(* Traffic lights *)', True),
    ('         lightSideA : BOOL;', False),
    ('\tlightSideB : BOOL;', False),
    ('END_VAR', True),
    ('PROCESS MainControl ', True),
    ('  VAR CONSTANT', True),
    ('    ROTATION_TIME : INT := 50;', False),
    ('  END_VAR', True),
    ('  VAR', True),
    ('    timer : INT := 0;', False),
    ('    manualActive : BOOL := FALSE;', False),
    ('    manualDirectionForward : BOOL := TRUE;', False),
    ('  END_VAR', True),
    ('  STATE idle', False),
    ('    IF rotateForwardButton AND NOT rotateBackwardButton THEN', False),
    ('      manualActive := TRUE;', False),
    ('      manualDirectionForward := TRUE;', False),
    ('      START PROCESS StartRotationForward;', False),
    ('      SET STATE manualRotation;', False),
    ('    ELSIF rotateBackwardButton AND NOT rotateForwardButton THEN', False),
    ('      manualActive := TRUE;', False),
    ('      manualDirectionForward := FALSE;', False),
    ('      START PROCESS StartRotationBackward;', False),
    ('      SET STATE manualRotation;', False),
    ('    ELSIF (presenceSideA OR presenceSideB) THEN', False),
    ('      manualActive := FALSE;', False),
    ('      manualDirectionForward := TRUE;', False),
    ('      START PROCESS StartRotationForward;', False),
    ('      timer := ROTATION_TIME;', False),
    ('      SET STATE rotating;', False),
    ('    END_IF', False),
    ('  END_STATE', False),
    ('  STATE rotating LOOPED', False),
    ('    IF (presenceSideA OR presenceSideB) THEN', False),
    ('      timer := ROTATION_TIME;', False),
    ('    END_IF', False),
    ('    IF partitionPressure THEN', False),
    ('      START PROCESS PauseRotation;', False),
    ('      SET STATE paused;', False),
    ('    END_IF', False),
    ('    timer := timer - 1;', False),
    ('    IF timer <= 0 THEN', False),
    ('      START PROCESS StopRotation;', False),
    ('      SET STATE idle;', False),
    ('    END_IF', False),
    ('  END_STATE', False),
    ('  STATE manualRotation LOOPED', False),
    ('    IF partitionPressure THEN', False),
    ('      START PROCESS PauseRotation;', False),
    ('      SET STATE paused;', False),
    ('    ELSIF rotateForwardButton AND rotateBackwardButton THEN', False),
    ('      START PROCESS StopRotation;', False),
    ('      manualActive := FALSE;', False),
    ('      SET STATE idle;', False),
    ('    ELSIF manualDirectionForward AND rotateForwardButton THEN', False),
    ('      START PROCESS StartRotationForward;', False),
    ('    ELSIF (NOT manualDirectionForward) AND rotateBackwardButton THEN', False),
    ('      START PROCESS StartRotationBackward;', False),
    ('    ELSE', False),
    ('      START PROCESS StopRotation;', False),
    ('      manualActive := FALSE;', False),
    ('      SET STATE idle;', False),
    ('    END_IF', False),
    ('  END_STATE', False),
    ('  STATE paused', False),
    ('    IF (PROCESS PauseRotation IN STATE STOP) THEN', False),
    ('      IF manualActive THEN', False),
    ('        IF manualDirectionForward AND rotateForwardButton THEN', False),
    ('          START PROCESS StartRotationForward;', False),
    ('          SET STATE manualRotation;', False),
    ('        ELSIF (NOT manualDirectionForward) AND rotateBackwardButton THEN', False),
    ('          START PROCESS StartRotationBackward;', False),
    ('          SET STATE manualRotation;', False),
    ('        ELSE', False),
    ('          START PROCESS StopRotation;', False),
    ('          manualActive := FALSE;', False),
    ('          SET STATE idle;', False),
    ('        END_IF', False),
    ('      ELSIF timer > 0 THEN', False),
    ('        START PROCESS StartRotationForward;', False),
    ('        SET STATE rotating;', False),
    ('      ELSE', False),
    ('        START PROCESS StopRotation;', False),
    ('        SET STATE idle;', False),
    ('      END_IF', False),
    ('    END_IF', False),
    ('  END_STATE', False),
    ('END_PROCESS', False),
    ('', False),
    ('PROCESS StartRotationForward', True),
    ('  STATE init', False),
    ('    motorOn := TRUE;', False),
    ('    rotationForward := TRUE;', False),
    ('    lightSideA := TRUE;', False),
    ('    lightSideB := TRUE;', False),
    ('    STOP;', False),
    ('  END_STATE', False),
    ('END_PROCESS', False),
    ('', False),
    ('PROCESS StartRotationBackward', True),
    ('  STATE init', False),
    ('    motorOn := TRUE;', False),
    ('    rotationForward := FALSE;', False),
    ('    lightSideA := TRUE;', False),
    ('    lightSideB := TRUE;', False),
    ('    STOP;', False),
    ('  END_STATE', False),
    ('END_PROCESS', False),
    ('', False),
    ('PROCESS StopRotation', True),
    ('  STATE init', False),
    ('    motorOn := FALSE;', False),
    ('    rotationForward := TRUE;', False),
    ('    lightSideA := TRUE;', False),
    ('    lightSideB := TRUE;', False),
    ('    STOP;', False),
    ('  END_STATE', False),
    ('END_PROCESS', False),
    ('', False),
    ('PROCESS PauseRotation', True),
    ('  VAR CONSTANT', True),
    ('    PAUSE_DURATION : INT := 30;', False),
    ('  END_VAR', True),
    ('  VAR', True),
    ('    counter : INT := 0;', False),
    ('  END_VAR', True),
    ('  STATE init', False),
    ('    motorOn := FALSE;', False),
    ('    lightSideA := FALSE;', False),
    ('    lightSideB := FALSE;', False),
    ('    counter := 0;', False),
    ('    SET NEXT;', False),
    ('  END_STATE', False),
    ('  STATE waiting LOOPED', False),
    ('    counter := counter + 1;', False),
    ('    IF counter >= PAUSE_DURATION THEN', False),
    ('      STOP;', False),
    ('    END_IF', False),
    ('  END_STATE', False),
    ('END_PROCESS', False),
    ('', False),
    ('', False),
    ('', False),
    ('', False),
    ('END_PROGRAM', True),
]

for text, bold in ctrl_lines:
    add_code_line(doc, text, bold=bold)

add_empty(doc)
add_empty(doc)
add_empty(doc)
add_empty(doc)

# ============ REQUIREMENTS ============
add_simple(doc, 'Требования', bold=True)

add_empty(doc)
add_empty(doc)
add_empty(doc)
add_empty(doc)

# ============ REQUIREMENTS TABLE ============
table = doc.add_table(rows=16, cols=7)
table.style = 'Table Grid'

# Header row - all bold, 10pt
headers = [
    'Description. NL requirement', 'trigger\nevent', 'release\nevent',
    'final\nevent', 'allowable\ndelay', 'invariant\ncondition', 'reaction'
]
for ci, h in enumerate(headers):
    add_table_cell_text(table.rows[0].cells[ci], h, bold=True)

# Data rows
rows_data = [
    # Row 1: default
    {0: [('default', False)],
     1: [('TRUE', True)], 2: [('FALSE', True)], 3: [('TRUE', True)],
     4: [('TRUE', True)], 5: [('TRUE', True)], 6: [('TRUE', True)]},
    # Row 2
    {0: [('Дверь должна начать вращаться при активации одного из двух сенсоров движения', False)],
     1: [('presenceSideA.RE OR presenceSideB.RE', False)],
     2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('motorOn', False)], 6: [('TRUE', True)]},
    # Row 3
    {0: [('Дверь должна вращаться в течение заданного времени ROTATION_TIME после последней активации сенсора', False)],
     1: [('presenceSideA.RE OR presenceSideB.RE', False)],
     2: [('FALSE', True)], 3: [('tau(#ROTATION_TIME)', False)], 4: [('TRUE', True)],
     5: [('motorOn', False)], 6: [('NOT ', True), ('motorOn', False)]},
    # Row 4
    {0: [('При давлении на перегородку вращение должно быть приостановлено', False)],
     1: [('partitionPressure', False)],
     2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('NOT ', True), ('motorOn', False)], 6: [('TRUE', True)]},
    # Row 5
    {0: [('После снятия давления вращение возобновляется через PAUSE_DURATION если таймер ещё не истёк', False)],
     1: [('partitionPressure.FE AND timer > 0', False)],
     2: [('FALSE', True)], 3: [('tau(#PAUSE_DURATION)', False)], 4: [('TRUE', True)],
     5: [('TRUE', True)], 6: [('motorOn', False)]},
    # Row 6
    {0: [('Индикатор на стороне A должен гореть зелёным при нормальном вращении', False)],
     1: [('motorOn AND ', False), ('NOT ', True), ('partitionPressure', False)],
     2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('lightSideA', False)], 6: [('TRUE', True)]},
    # Row 7
    {0: [('Индикатор на стороне B должен гореть зелёным при нормальном вращении', False)],
     1: [('motorOn AND ', False), ('NOT ', True), ('partitionPressure', False)],
     2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('lightSideB', False)], 6: [('TRUE', True)]},
    # Row 8
    {0: [('Индикаторы должны гореть красным при паузе (давление на перегородку)', False)],
     1: [('partitionPressure', False)],
     2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('NOT ', True), ('lightSideA AND ', False), ('NOT ', True), ('lightSideB', False)],
     6: [('TRUE', True)]},
    # Row 9
    {0: [('Дверь не должна вращаться при истёкшем таймере и отсутствии активации сенсоров', False)],
     1: [('TRUE', True)], 2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('NOT ', True), ('(motorOn AND timer <= 0)', False)],
     6: [('TRUE', True)]},
    # Row 10
    {0: [('Индикаторы должны гореть зелёным в исходном состоянии (дверь стоит, вход разрешён)', False)],
     1: [('NOT ', True), ('motorOn AND ', False), ('NOT ', True), ('partitionPressure', False)],
     2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('lightSideA AND lightSideB', False)],
     6: [('TRUE', True)]},
    # Row 11 - empty
    {0: [('', False)], 1: [('', False)], 2: [('', False)], 3: [('', False)],
     4: [('', False)], 5: [('', False)], 6: [('', False)]},
]

rows_data = [
    {0: [('default', False)],
     1: [('TRUE', True)], 2: [('FALSE', True)], 3: [('TRUE', True)],
     4: [('TRUE', True)], 5: [('TRUE', True)], 6: [('TRUE', True)]},
    {0: [('Дверь должна начать вращаться при активации одного из двух сенсоров движения', False)],
     1: [('presenceSideA.RE OR presenceSideB.RE', False)],
     2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('motorOn AND rotationForward', False)], 6: [('TRUE', True)]},
    {0: [('Дверь должна вращаться в течение заданного времени ROTATION_TIME после последней активации сенсора', False)],
     1: [('presenceSideA.RE OR presenceSideB.RE', False)],
     2: [('FALSE', True)], 3: [('tau(#ROTATION_TIME)', False)], 4: [('TRUE', True)],
     5: [('motorOn AND rotationForward', False)], 6: [('NOT ', True), ('motorOn', False)]},
    {0: [('При нажатии rotateForwardButton дверь должна начать вращаться вперёд', False)],
     1: [('rotateForwardButton.RE', False)],
     2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('motorOn AND rotationForward', False)], 6: [('TRUE', True)]},
    {0: [('При нажатии rotateBackwardButton дверь должна начать вращаться назад', False)],
     1: [('rotateBackwardButton.RE', False)],
     2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('motorOn AND NOT rotationForward', False)], 6: [('TRUE', True)]},
    {0: [('Одновременное нажатие rotateForwardButton и rotateBackwardButton не должно запускать двигатель', False)],
     1: [('rotateForwardButton AND rotateBackwardButton', False)],
     2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('NOT motorOn', False)], 6: [('TRUE', True)]},
    {0: [('При давлении на перегородку вращение должно быть приостановлено', False)],
     1: [('partitionPressure', False)],
     2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('NOT ', True), ('motorOn', False)], 6: [('TRUE', True)]},
    {0: [('После снятия давления вращение возобновляется через PAUSE_DURATION если таймер ещё не истёк', False)],
     1: [('partitionPressure.FE AND timer > 0', False)],
     2: [('FALSE', True)], 3: [('tau(#PAUSE_DURATION)', False)], 4: [('TRUE', True)],
     5: [('TRUE', True)], 6: [('motorOn', False)]},
    {0: [('Индикатор на стороне A должен гореть зелёным при нормальном вращении', False)],
     1: [('motorOn AND ', False), ('NOT ', True), ('partitionPressure', False)],
     2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('lightSideA', False)], 6: [('TRUE', True)]},
    {0: [('Индикатор на стороне B должен гореть зелёным при нормальном вращении', False)],
     1: [('motorOn AND ', False), ('NOT ', True), ('partitionPressure', False)],
     2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('lightSideB', False)], 6: [('TRUE', True)]},
    {0: [('Индикаторы должны гореть красным при паузе (давление на перегородку)', False)],
     1: [('partitionPressure', False)],
     2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('NOT ', True), ('lightSideA AND ', False), ('NOT ', True), ('lightSideB', False)],
     6: [('TRUE', True)]},
    {0: [('Дверь не должна вращаться при истёкшем таймере и отсутствии активации сенсоров', False)],
     1: [('TRUE', True)], 2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('NOT ', True), ('(motorOn AND timer <= 0)', False)],
     6: [('TRUE', True)]},
    {0: [('Индикаторы должны гореть зелёным в исходном состоянии (дверь стоит, вход разрешён)', False)],
     1: [('NOT ', True), ('motorOn AND ', False), ('NOT ', True), ('partitionPressure', False)],
     2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('lightSideA AND lightSideB', False)],
     6: [('TRUE', True)]},
    {0: [('При включённом двигателе датчик isRotating должен быть активен', False)],
     1: [('motorOn', False)], 2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('isRotating', False)], 6: [('TRUE', True)]},
    {0: [('При вращении значение doorAngle должно изменяться', False)],
     1: [('motorOn', False)], 2: [('FALSE', True)], 3: [('TRUE', True)], 4: [('TRUE', True)],
     5: [('doorAngle >= 0.0', False)], 6: [('TRUE', True)]},
]

for ri, row_dict in enumerate(rows_data):
    for ci, runs in row_dict.items():
        cell = table.rows[ri + 1].cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        for text, bold in runs:
            r = p.add_run(text)
            r.font.size = SZ10
            if bold:
                r.font.bold = True

add_empty(doc)
add_empty(doc)

add_simple(doc, 'Алгоритм управления', bold=True)

add_empty(doc)

add_simple(doc, 'Оформление', bold=True)

add_empty(doc)

# ============ SAVE ============
doc.save('RevolvingDoor_Sample.docx')
print('OK: RevolvingDoor_Sample.docx created successfully')
