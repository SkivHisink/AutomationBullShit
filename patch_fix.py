"""
Fix the VAR block ordering and missing timer reset in MainControl.
"""
import shutil, copy
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

DST = r'f:\Github\AutomationBullShit\RevolvingDoor_Sample.docx'
doc = Document(DST)

def find_para_idx(substr, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if substr in p.text:
            return i
    return None

def delete_para_at(idx):
    elem = doc.paragraphs[idx]._element
    elem.getparent().remove(elem)

def clone_with_text(template_idx, text, bold=False):
    """Clone paragraph at template_idx, set text, return new XML element."""
    tmpl = doc.paragraphs[template_idx]._element
    new_elem = copy.deepcopy(tmpl)
    for r in new_elem.findall(qn('w:r')):
        for t in r.findall(qn('w:t')):
            t.text = ''
    first_r = new_elem.find(qn('w:r'))
    if first_r is not None:
        t_elem = first_r.find(qn('w:t'))
        if t_elem is not None:
            t_elem.text = text
            t_elem.set(qn('xml:space'), 'preserve')
        rpr = first_r.find(qn('w:rPr'))
        if rpr is not None:
            b_elem = rpr.find(qn('w:b'))
            if bold and b_elem is None:
                etree.SubElement(rpr, qn('w:b'))
            elif not bold and b_elem is not None:
                rpr.remove(b_elem)
    return new_elem

# ── 1. Fix VAR block order in MainControl ────────────────────────────
# Current wrong order (P157-P163):
#   PROCESS MainControl
#   END_VAR          ← wrong
#   timer : INT := 0;
#   VAR
#   END_VAR
#   ROTATION_TIME : INT := 50;
#   VAR CONSTANT     ← wrong
#   STATE idle

# Need to remove P158-P163 and re-insert in correct order:
#   VAR CONSTANT
#   ROTATION_TIME : INT := 50;
#   END_VAR
#   VAR
#   timer : INT := 0;
#   END_VAR

mc_idx = find_para_idx('PROCESS MainControl')
print(f'MainControl at P{mc_idx}')

# Verify the wrong order
for i in range(mc_idx, mc_idx + 10):
    print(f'  P{i}: {doc.paragraphs[i].text.strip()[:60]}')

# Delete the 6 wrongly-ordered paragraphs (P158-P163)
# Delete from bottom to top to avoid index shifting
idle_idx = find_para_idx('STATE idle', mc_idx)
print(f'STATE idle at P{idle_idx}')

# Remove all paragraphs between PROCESS MainControl and STATE idle
to_remove = []
for i in range(mc_idx + 1, idle_idx):
    to_remove.append(doc.paragraphs[i]._element)

for elem in to_remove:
    elem.getparent().remove(elem)

# Now re-find STATE idle (index shifted)
idle_idx = find_para_idx('STATE idle', mc_idx)
print(f'STATE idle now at P{idle_idx}')

# Insert correct VAR blocks before STATE idle
idle_elem = doc.paragraphs[idle_idx]._element
correct_lines = [
    ('  VAR CONSTANT', True),
    ('    ROTATION_TIME : INT := 50;', False),
    ('  END_VAR', True),
    ('  VAR', True),
    ('    timer : INT := 0;', False),
    ('  END_VAR', True),
]

# Insert in forward order, each before idle
prev = idle_elem
for text, bold in reversed(correct_lines):
    new_elem = clone_with_text(idle_idx, text, bold)
    idle_elem.addprevious(new_elem)

# ── 2. Fix "SET STATE rotating LOOPED;" → "SET STATE rotating;" ─────
for p in doc.paragraphs:
    if p.text.strip() == 'SET STATE rotating LOOPED;':
        for r in p.runs:
            r.text = r.text.replace('SET STATE rotating LOOPED;', 'SET STATE rotating;')
            if 'LOOPED' in r.text:
                r.text = r.text.replace(' LOOPED', '')

# ── 3. Fix "STATE rotating" → "STATE rotating LOOPED" ───────────────
for p in doc.paragraphs:
    if p.text.strip() == 'STATE rotating':
        runs = p.runs
        if runs:
            # Find the run with "rotating"
            for r in runs:
                if 'rotating' in r.text and 'LOOPED' not in r.text:
                    r.text = r.text.replace('rotating', 'rotating LOOPED')
                    break

# ── 4. Add timer reset in rotating state ─────────────────────────────
# Need to add before "IF partitionPressure THEN" in rotating state:
#   IF (presenceSideA OR presenceSideB) THEN
#     timer := ROTATION_TIME;
#   END_IF

# Find rotating LOOPED state, then find the IF partitionPressure inside it
rot_idx = find_para_idx('STATE rotating LOOPED')
if rot_idx is not None:
    part_idx = find_para_idx('IF partitionPressure THEN', rot_idx)
    if part_idx is not None:
        part_elem = doc.paragraphs[part_idx]._element
        timer_lines = [
            ('    IF (presenceSideA OR presenceSideB) THEN', False),
            ('      timer := ROTATION_TIME;', False),
            ('    END_IF', False),
        ]
        for text, bold in reversed(timer_lines):
            new_elem = clone_with_text(part_idx, text, bold)
            part_elem.addprevious(new_elem)
        print('Added timer reset block in rotating state')

# ── Verify ───────────────────────────────────────────────────────────
print('\n=== MainControl section after fix ===')
mc_idx = find_para_idx('PROCESS MainControl')
end_idx = find_para_idx('END_PROCESS', mc_idx)
if end_idx:
    for i in range(mc_idx, end_idx + 1):
        print(f'  P{i}: {doc.paragraphs[i].text.strip()[:80]}')

doc.save(DST)
print('\nSaved!')
