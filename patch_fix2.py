"""
Fix MainControl section by working directly with XML elements.
Start fresh from the good file.
"""
import shutil, copy
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

SRC = r'F:\Github\RevolvingDoor_Sample.docx'
DST = r'f:\Github\AutomationBullShit\RevolvingDoor_Sample.docx'

# Start fresh from good file
shutil.copy2(SRC, DST)
doc = Document(DST)
body = doc.element.body

# Get all paragraph elements
all_paras = body.findall(qn('w:p'))

def get_text(p_elem):
    texts = []
    for r in p_elem.findall(qn('w:r')):
        for t in r.findall(qn('w:t')):
            if t.text:
                texts.append(t.text)
    return ''.join(texts).strip()

def find_para_elem(substr, start_idx=0):
    for i, p in enumerate(all_paras):
        if i < start_idx:
            continue
        if substr in get_text(p):
            return i, p
    return None, None

def make_new_para(template, text, bold=None):
    """Create new paragraph element from template with given text."""
    new = copy.deepcopy(template)
    # Clear all runs
    for r in new.findall(qn('w:r')):
        new.remove(r)
    # Clone first run from template
    orig_runs = template.findall(qn('w:r'))
    if orig_runs:
        new_r = copy.deepcopy(orig_runs[0])
        for t in new_r.findall(qn('w:t')):
            t.text = text
            t.set(qn('xml:space'), 'preserve')
        if bold is not None:
            rpr = new_r.find(qn('w:rPr'))
            if rpr is not None:
                b = rpr.find(qn('w:b'))
                if bold and b is None:
                    etree.SubElement(rpr, qn('w:b'))
                elif not bold and b is not None:
                    rpr.remove(b)
        new.append(new_r)
    return new

def set_para_text(p_elem, text):
    """Replace all text in paragraph, keep first run formatting."""
    runs = p_elem.findall(qn('w:r'))
    if runs:
        # Set first run text, remove others
        for t in runs[0].findall(qn('w:t')):
            t.text = text
            t.set(qn('xml:space'), 'preserve')
        for r in runs[1:]:
            p_elem.remove(r)

def rebuild_runs(p_elem, run_specs, size_emu=177800):
    """Replace all runs with new ones. run_specs = [(text, bold), ...]"""
    from docx.oxml import OxmlElement
    # Remove all runs
    for r in p_elem.findall(qn('w:r')):
        p_elem.remove(r)
    # Add new runs
    for text, bold in run_specs:
        r = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size_emu // 6350))  # EMU to half-points
        rpr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(size_emu // 6350))
        rpr.append(szCs)
        if bold:
            b = OxmlElement('w:b')
            rpr.append(b)
        r.append(rpr)
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p_elem.append(r)

# ════════════════════════════════════════════════════════════════════
# Apply all text changes
# ════════════════════════════════════════════════════════════════════

# Refresh para list
all_paras = list(body.findall(qn('w:p')))

# ── 1. Title: "Вращающаяся дверь" → "Вращающаяся дверь. " ───────────
i, p = find_para_elem('Вращающаяся дверь')
if p is not None:
    txt = get_text(p)
    if not txt.endswith('.'):
        runs = p.findall(qn('w:r'))
        if runs:
            t_elems = runs[-1].findall(qn('w:t'))
            if t_elems:
                t_elems[-1].text = t_elems[-1].text.rstrip() + '. '

# ── 2. Add "секция – Section (A, B, C)" after "Описание задачи." ────
i_desc, p_desc = find_para_elem('Описание задачи.')
if p_desc is not None:
    # Insert empty para + section line after Описание задачи
    next_sib = p_desc.getnext()
    # Create empty paragraph
    empty_p = copy.deepcopy(p_desc)
    for r in empty_p.findall(qn('w:r')):
        empty_p.remove(r)
    # Create section line paragraph
    sec_p = make_new_para(p_desc, 'секция – Section (A, B, C)', bold=False)

    if next_sib is not None:
        next_sib.addprevious(empty_p)
        next_sib.addprevious(sec_p)
    else:
        body.append(empty_p)
        body.append(sec_p)

# Refresh
all_paras = list(body.findall(qn('w:p')))

# ── 3. Main description paragraph ───────────────────────────────────
i, p = find_para_elem('Вращающаяся дверь — это трёхсекционная')
if p is not None:
    rebuild_runs(p, [
        ('Вращающаяся дверь — это трёхсекционная (', False),
        ('SECTION_COUNT', True),
        (' = 3) дверь, установленная в проёме здания, обеспечивающая проход людей с двух сторон (', False),
        ('SideA / SideB', True),
        ('). В нормальном (исходном) состоянии дверь неподвижна. '
         'При приближении пользователя к одной из сторон (датчики ', False),
        ('presenceSideA / presenceSideB', True),
        (') двигатель (', False),
        ('motorOn', True),
        (') запускает вращение двери с заданной скоростью (', False),
        ('ROTATION_SPEED', True),
        ('). Вращение продолжается в течение заданного времени (', False),
        ('ROTATION_TIME', True),
        ('). Если за время вращения повторно срабатывает один из сенсоров, '
         'таймер сбрасывается и отсчёт начинается заново. '
         'Вращение останавливается по истечении таймера. '
         'При оказании давления на секционные перегородки (датчик ', False),
        ('partitionPressure', True),
        (') вращение приостанавливается на время ', False),
        ('PAUSE_DURATION', True),
        (', после чего возобновляется, если таймер ещё не истёк.', False),
    ])
    print('Updated main description')

# ── 4. Speed paragraph ──────────────────────────────────────────────
i, p = find_para_elem('Скорость вращения двери регулируется')
if p is not None:
    rebuild_runs(p, [
        ('Дверь вращается с постоянной скоростью (', False),
        ('ROTATION_SPEED', True),
        ('). Время вращения после последнего срабатывания сенсора задаётся константой ', False),
        ('ROTATION_TIME', True),
        ('.', False),
    ])
    print('Updated speed paragraph')

# ── 5. Remove "нахождение пользователя в зоне вращения" line ────────
all_paras = list(body.findall(qn('w:p')))
i, p = find_para_elem('Оператор может имитировать нахождение пользователя в зоне вращения')
if p is not None:
    body.remove(p)
    print('Removed userInDoor operator line')

# ── 6. Controller description lines ─────────────────────────────────
all_paras = list(body.findall(qn('w:p')))
i, p = find_para_elem('Двигатель запускается при обнаружении пользователя')
if p is not None:
    set_para_text(p, 'Двигатель запускается при срабатывании одного из двух сенсоров движения.')
    print('Updated motor start description')

all_paras = list(body.findall(qn('w:p')))
i, p = find_para_elem('Двигатель останавливается, когда зона вращения пуста')
if p is not None:
    set_para_text(p, 'Двигатель останавливается по истечении таймера ROTATION_TIME.')
    print('Updated motor stop description')

# ── 7. Environment: remove EnterDoor, LeaveDoor ─────────────────────
all_paras = list(body.findall(qn('w:p')))
i, p = find_para_elem('Environment:')
if p is not None:
    full_text = get_text(p)
    new_text = full_text.replace('EnterDoor, LeaveDoor, ', '').replace(', EnterDoor, LeaveDoor', '')
    set_para_text(p, new_text)
    print('Updated Environment line')

# ── 8. Constants: SAFE_SPEED → ROTATION_TIME (all occurrences) ──────
all_paras = list(body.findall(qn('w:p')))
for p in all_paras:
    txt = get_text(p)
    if 'SAFE_SPEED' in txt:
        # Replace in all text runs
        for r in p.findall(qn('w:r')):
            for t in r.findall(qn('w:t')):
                if t.text and 'SAFE_SPEED' in t.text:
                    t.text = t.text.replace('SAFE_SPEED', 'ROTATION_TIME')
        print(f'Replaced SAFE_SPEED in: {txt[:60]}')

# ── 9. Controls: remove userInDoor ───────────────────────────────────
all_paras = list(body.findall(qn('w:p')))
for p in all_paras:
    txt = get_text(p)
    if 'userInDoor' in txt and 'Controls' in txt:
        new_text = txt.replace(', userInDoor', '').replace('userInDoor, ', '')
        set_para_text(p, new_text)
        print(f'Removed userInDoor from Controls')

# ── 10. Remove "Имитировать вход пользователя в зону вращения" ──────
all_paras = list(body.findall(qn('w:p')))
i, p = find_para_elem('Имитировать вход пользователя в зону вращения')
if p is not None:
    body.remove(p)
    print('Removed enter door line')

# ── 11. Change Controls "Имитировать вход/выход" ────────────────────
all_paras = list(body.findall(qn('w:p')))
i, p = find_para_elem('Имитировать вход/выход пользователя')
if p is not None:
    set_para_text(p, 'При срабатывании сенсора дверь вращается ROTATION_TIME. При закрытых индикаторах — ошибка.')
    print('Updated Controls enter/exit line')

# ── 12. VAR_GLOBAL: remove userInDoor lines ─────────────────────────
all_paras = list(body.findall(qn('w:p')))
to_remove = []
for p in all_paras:
    txt = get_text(p)
    if txt == '(* User in rotation zone *)':
        to_remove.append(p)
    elif txt in ('userInDoor  : BOOL;', 'userInDoor : BOOL;'):
        to_remove.append(p)
for p in to_remove:
    body.remove(p)
    print('Removed userInDoor VAR line')

# ── 13. Plant Init: remove "userInDoor := FALSE;" ───────────────────
all_paras = list(body.findall(qn('w:p')))
for p in all_paras:
    if get_text(p) == 'userInDoor := FALSE;':
        body.remove(p)
        print('Removed userInDoor init')

# ── 14. Controller VAR: remove userInDoor ────────────────────────────
# Find PROGRAM Controller, then remove userInDoor lines after it
all_paras = list(body.findall(qn('w:p')))
ctrl_found = False
ctrl_to_remove = []
for p in all_paras:
    txt = get_text(p)
    if txt == 'PROGRAM Controller':
        ctrl_found = True
    if ctrl_found:
        if txt == '(* User in rotation zone *)':
            ctrl_to_remove.append(p)
        elif txt in ('userInDoor  : BOOL;', 'userInDoor : BOOL;'):
            ctrl_to_remove.append(p)
        if txt == 'END_PROGRAM':
            break
for p in ctrl_to_remove:
    body.remove(p)
    print('Removed userInDoor from Controller VAR')

# ── 15. Controller MainControl: add timer logic ─────────────────────
all_paras = list(body.findall(qn('w:p')))

# Find "PROCESS MainControl" and "STATE idle" after it
mc_idx, mc_p = find_para_elem('PROCESS MainControl')
idle_idx, idle_p = find_para_elem('STATE idle', mc_idx)

print(f'MainControl at idx {mc_idx}, STATE idle at idx {idle_idx}')

# Insert VAR CONSTANT / ROTATION_TIME / END_VAR / VAR / timer / END_VAR before STATE idle
var_lines = [
    ('  VAR CONSTANT', True),
    ('    ROTATION_TIME : INT := 50;', False),
    ('  END_VAR', True),
    ('  VAR', True),
    ('    timer : INT := 0;', False),
    ('  END_VAR', True),
]

# Insert each BEFORE idle_p, in forward order
for text, bold in var_lines:
    new_p = make_new_para(idle_p, text, bold)
    idle_p.addprevious(new_p)
print('Inserted VAR blocks')

# ── 16. Add "timer := ROTATION_TIME;" after "START PROCESS StartRotation;" in idle state
all_paras = list(body.findall(qn('w:p')))
# Find first "START PROCESS StartRotation;" after MainControl
for i, p in enumerate(all_paras):
    if get_text(p) == 'START PROCESS StartRotation;' and i > mc_idx:
        # Check context: next should be "SET STATE rotating;"
        next_p = all_paras[i + 1] if i + 1 < len(all_paras) else None
        if next_p is not None and 'SET STATE rotating' in get_text(next_p):
            timer_p = make_new_para(p, '      timer := ROTATION_TIME;', False)
            next_p.addprevious(timer_p)
            print('Added timer := ROTATION_TIME in idle state')
            break

# ── 17. "STATE rotating" → "STATE rotating LOOPED" ──────────────────
all_paras = list(body.findall(qn('w:p')))
for p in all_paras:
    txt = get_text(p)
    if txt == 'STATE rotating':
        for r in p.findall(qn('w:r')):
            for t in r.findall(qn('w:t')):
                if t.text and 'rotating' in t.text:
                    t.text = t.text.replace('rotating', 'rotating LOOPED')
        print('Changed STATE rotating → STATE rotating LOOPED')
        break

# ── 18. Add timer reset + timer decrement in rotating state ──────────
# Current rotating state:
#   STATE rotating LOOPED
#   IF partitionPressure THEN ... END_IF
#   IF (NOT userInDoor ...) THEN ... StopRotation ... idle ... END_IF
#   END_STATE
#
# Need:
#   STATE rotating LOOPED
#   IF (presenceSideA OR presenceSideB) THEN
#     timer := ROTATION_TIME;
#   END_IF
#   IF partitionPressure THEN ... END_IF
#   timer := timer - 1;
#   IF timer <= 0 THEN
#     START PROCESS StopRotation;
#     SET STATE idle;
#   END_IF
#   END_STATE

all_paras = list(body.findall(qn('w:p')))
rot_idx, rot_p = find_para_elem('STATE rotating LOOPED')
print(f'STATE rotating LOOPED at idx {rot_idx}')

# Find partitionPressure check in rotating state
part_idx, part_p = find_para_elem('IF partitionPressure THEN', rot_idx)

# Insert timer reset before partitionPressure
timer_reset = [
    ('    IF (presenceSideA OR presenceSideB) THEN', False),
    ('      timer := ROTATION_TIME;', False),
    ('    END_IF', False),
]
for text, bold in timer_reset:
    new_p = make_new_para(part_p, text, bold)
    part_p.addprevious(new_p)
print('Added timer reset in rotating state')

# Find and replace "IF (NOT userInDoor..." block with timer logic
all_paras = list(body.findall(qn('w:p')))
for i, p in enumerate(all_paras):
    txt = get_text(p)
    if 'NOT userInDoor AND NOT presenceSideA AND NOT presenceSideB' in txt:
        # This is the old stop condition. Replace this and the following lines up to END_IF
        # Replace this line with timer decrement + timer check
        # First, find the END_IF that closes this block
        # Lines: IF (...) THEN / START PROCESS StopRotation; / SET STATE idle; / END_IF
        set_para_text(p, '    timer := timer - 1;')

        # Insert "IF timer <= 0 THEN" after timer decrement
        next_p = all_paras[i + 1]
        timer_check = make_new_para(p, '    IF timer <= 0 THEN', False)
        next_p.addprevious(timer_check)
        print('Replaced userInDoor check with timer logic')
        break

# ── 19. Replace "IF userInDoor THEN" → "IF timer > 0 THEN" in paused state
all_paras = list(body.findall(qn('w:p')))
for p in all_paras:
    txt = get_text(p)
    if txt == 'IF userInDoor THEN':
        set_para_text(p, '      IF timer > 0 THEN')
        print('Replaced userInDoor check in paused state')
        break

# ── 20. Table changes ───────────────────────────────────────────────
from docx.shared import Pt
SZ10 = Pt(10)

def set_cell_runs(cell, run_specs):
    """Set cell text. run_specs = [(text, bold), ...]"""
    para = cell.paragraphs[0]
    for r_elem in para._element.findall(qn('w:r')):
        para._element.remove(r_elem)
    for text, bold in run_specs:
        r = para.add_run(text)
        r.font.size = SZ10
        if bold:
            r.font.bold = True

table = doc.tables[0]

# Row 2: sensor activation
set_cell_runs(table.cell(2, 0), [('Дверь должна начать вращаться при активации одного из двух сенсоров движения', False)])
set_cell_runs(table.cell(2, 1), [('presenceSideA.RE OR presenceSideB.RE', False)])

# Row 3: timer-based rotation
set_cell_runs(table.cell(3, 0), [('Дверь должна вращаться в течение заданного времени ROTATION_TIME после последней активации сенсора', False)])
set_cell_runs(table.cell(3, 1), [('presenceSideA.RE OR presenceSideB.RE', False)])
set_cell_runs(table.cell(3, 3), [('tau(#ROTATION_TIME)', False)])
set_cell_runs(table.cell(3, 5), [('motorOn', False)])
set_cell_runs(table.cell(3, 6), [('NOT ', True), ('motorOn', False)])

# Row 5: timer condition
set_cell_runs(table.cell(5, 0), [('После снятия давления вращение возобновляется через PAUSE_DURATION если таймер ещё не истёк', False)])
set_cell_runs(table.cell(5, 1), [('partitionPressure.FE AND timer > 0', False)])

# Row 9: timer-based invariant
set_cell_runs(table.cell(9, 0), [('Дверь не должна вращаться при истёкшем таймере и отсутствии активации сенсоров', False)])
set_cell_runs(table.cell(9, 5), [('NOT ', True), ('(motorOn AND timer <= 0)', False)])

doc.save(DST)
print('\n=== DONE ===')

# Verify Controller section
doc2 = Document(DST)
in_ctrl = False
for p in doc2.paragraphs:
    txt = p.text.strip()
    if txt == 'PROCESS MainControl':
        in_ctrl = True
    if in_ctrl:
        print(f'  {txt}')
    if in_ctrl and txt == 'END_PROCESS':
        in_ctrl = False
        break
